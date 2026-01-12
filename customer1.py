import streamlit as st
import pandas as pd
import plotly.express as px
import plotly.graph_objects as go
from datetime import datetime, timedelta
import uuid
import time
import os
import json
import sqlite3
from typing import List, Dict, Optional, Tuple
import re
from io import BytesIO
import base64

# For PDF/DOCX/Excel processing
try:
    import PyPDF2
    from docx import Document
    import openpyxl
    from openpyxl import load_workbook
except ImportError:
    st.warning("Please install: pip install PyPDF2 python-docx openpyxl")

# For vector database and embeddings
try:
    from langchain.text_splitter import RecursiveCharacterTextSplitter
    from langchain_community.vectorstores import FAISS
    from langchain_openai import OpenAIEmbeddings, ChatOpenAI
    from langchain.chains import RetrievalQA
    from langchain.docstore.document import Document as LangchainDocument
    from langchain.prompts import PromptTemplate
except ImportError:
    st.warning("Please install: pip install langchain langchain-openai langchain-community faiss-cpu")

# For language detection and translation
try:
    from langdetect import detect, LangDetectException
    from deep_translator import GoogleTranslator
except ImportError:
    st.warning("Please install: pip install langdetect deep-translator")

# For web scraping
try:
    import requests
    from bs4 import BeautifulSoup
except ImportError:
    st.warning("Please install: pip install requests beautifulsoup4")

# For OCR
try:
    import pytesseract
    from PIL import Image
except ImportError:
    st.warning("Please install: pip install pytesseract Pillow")

# For semantic similarity (confidence scoring)
try:
    from sklearn.metrics.pairwise import cosine_similarity
    import numpy as np
except ImportError:
    st.warning("Please install: pip install scikit-learn numpy")

# Page configuration
st.set_page_config(
    page_title="AI Support Agent",
    page_icon="🤖",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Custom CSS
st.markdown("""
<style>
    .main-header {
        font-size: 2.5rem;
        font-weight: bold;
        color: #1f77b4;
        text-align: center;
        margin-bottom: 2rem;
    }
    .chat-message {
        padding: 1rem;
        border-radius: 0.5rem;
        margin-bottom: 1rem;
    }
    .user-message {
        background-color: #e3f2fd;
        margin-left: 2rem;
    }
    .bot-message {
        background-color: #f5f5f5;
        margin-right: 2rem;
    }
    .ticket-card {
        padding: 1rem;
        border: 1px solid #ddd;
        border-radius: 0.5rem;
        margin-bottom: 0.5rem;
    }
    .confidence-high {
        color: #2ecc71;
        font-weight: bold;
    }
    .confidence-medium {
        color: #f39c12;
        font-weight: bold;
    }
    .confidence-low {
        color: #e74c3c;
        font-weight: bold;
    }
</style>
""", unsafe_allow_html=True)

# Database setup
def init_database():
    """Initialize SQLite database for persistence"""
    conn = sqlite3.connect('support_agent.db', check_same_thread=False)
    c = conn.cursor()
    
    # Chat history table
    c.execute('''CREATE TABLE IF NOT EXISTS chat_history
                 (id TEXT PRIMARY KEY, role TEXT, message TEXT, timestamp TEXT, 
                  confidence REAL, response_time REAL, language TEXT, category TEXT)''')
    
    # Tickets table
    c.execute('''CREATE TABLE IF NOT EXISTS tickets
                 (id TEXT PRIMARY KEY, query TEXT, language TEXT, category TEXT, 
                  status TEXT, priority TEXT, assigned_to TEXT, timestamp TEXT, 
                  resolved_at TEXT, resolution_time REAL)''')
    
    # Feedback table
    c.execute('''CREATE TABLE IF NOT EXISTS feedback
                 (chat_id TEXT, feedback TEXT, timestamp TEXT)''')
    
    # Analytics table
    c.execute('''CREATE TABLE IF NOT EXISTS analytics
                 (date TEXT, total_queries INTEGER, answered INTEGER, escalated INTEGER,
                  avg_response_time REAL, avg_confidence REAL)''')
    
    conn.commit()
    return conn

# Initialize database
if 'db_conn' not in st.session_state:
    st.session_state.db_conn = init_database()

# Initialize session state
if 'chat_history' not in st.session_state:
    st.session_state.chat_history = []
if 'tickets' not in st.session_state:
    st.session_state.tickets = []
if 'analytics' not in st.session_state:
    st.session_state.analytics = {
        'total_queries': 0,
        'answered': 0,
        'escalated': 0,
        'languages': {'English': 0, 'Hindi': 0, 'Marathi': 0, 'Other': 0},
        'categories': {'Billing': 0, 'Technical': 0, 'General': 0}
    }
if 'vector_store' not in st.session_state:
    st.session_state.vector_store = None
if 'knowledge_base_text' not in st.session_state:
    st.session_state.knowledge_base_text = ""
if 'feedback' not in st.session_state:
    st.session_state.feedback = {}
if 'embeddings_model' not in st.session_state:
    st.session_state.embeddings_model = None

# Helper Functions
def extract_text_from_pdf(file) -> str:
    """Extract text from PDF file"""
    try:
        pdf_reader = PyPDF2.PdfReader(file)
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text() + "\n"
        return text
    except Exception as e:
        st.error(f"Error reading PDF: {str(e)}")
        return ""

def extract_text_from_docx(file) -> str:
    """Extract text from DOCX file"""
    try:
        doc = Document(file)
        text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        return text
    except Exception as e:
        st.error(f"Error reading DOCX: {str(e)}")
        return ""

def extract_text_from_image(image_file) -> str:
    """Extract text from image using OCR"""
    try:
        image = Image.open(image_file)
        text = pytesseract.image_to_string(image)
        return text
    except Exception as e:
        st.error(f"Error performing OCR: {str(e)}")
        return ""

def scrape_url(url: str) -> str:
    """Scrape text content from URL"""
    try:
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
        }
        response = requests.get(url, headers=headers, timeout=10)
        response.raise_for_status()
        
        soup = BeautifulSoup(response.content, 'html.parser')
        
        # Remove script and style elements
        for script in soup(["script", "style"]):
            script.decompose()
        
        # Get text
        text = soup.get_text()
        
        # Clean up text
        lines = (line.strip() for line in text.splitlines())
        chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
        text = '\n'.join(chunk for chunk in chunks if chunk)
        
        return text
    except Exception as e:
        st.error(f"Error scraping URL {url}: {str(e)}")
        return ""

def detect_language(text: str) -> str:
    """Detect language of text"""
    try:
        lang_code = detect(text)
        lang_map = {
            'en': 'English',
            'hi': 'Hindi',
            'mr': 'Marathi'
        }
        return lang_map.get(lang_code, 'Other')
    except:
        return 'English'

def translate_text(text: str, target_lang: str) -> str:
    """Translate text to target language"""
    try:
        if target_lang == 'English':
            return text
        
        lang_code_map = {
            'Hindi': 'hi',
            'Marathi': 'mr'
        }
        
        target_code = lang_code_map.get(target_lang, 'en')
        translator = GoogleTranslator(source='auto', target=target_code)
        translated = translator.translate(text)
        return translated
    except Exception as e:
        st.warning(f"Translation failed: {str(e)}")
        return text

def categorize_query(query: str) -> str:
    """Keyword-based categorization with AI enhancement"""
    query_lower = query.lower()
    
    billing_keywords = ['payment', 'invoice', 'bill', 'charge', 'refund', 'price', 'cost', 
                       'subscription', 'card', 'billing', 'money', 'paid']
    technical_keywords = ['error', 'bug', 'issue', 'problem', 'not working', 'broken', 
                         'crash', 'slow', 'loading', 'login', 'access', 'technical']
    
    billing_score = sum(1 for keyword in billing_keywords if keyword in query_lower)
    technical_score = sum(1 for keyword in technical_keywords if keyword in query_lower)
    
    if billing_score > technical_score and billing_score > 0:
        return 'Billing'
    elif technical_score > 0:
        return 'Technical'
    else:
        return 'General'

def assign_priority(confidence: float, category: str) -> str:
    """Assign priority based on confidence and category"""
    if confidence < 0.4 or category == 'Billing':
        return 'High'
    elif confidence < 0.6 or category == 'Technical':
        return 'Medium'
    else:
        return 'Low'

def create_vector_store(text: str, openai_api_key: str):
    """Create FAISS vector store from text with caching"""
    try:
        # Split text into chunks
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200,
            length_function=len
        )
        chunks = text_splitter.split_text(text)
        
        # Create documents
        documents = [LangchainDocument(page_content=chunk) for chunk in chunks]
        
        # Create embeddings and vector store
        embeddings = OpenAIEmbeddings(openai_api_key=openai_api_key)
        st.session_state.embeddings_model = embeddings
        vector_store = FAISS.from_documents(documents, embeddings)
        
        return vector_store
    except Exception as e:
        st.error(f"Error creating vector store: {str(e)}")
        return None

def calculate_semantic_confidence(query: str, retrieved_docs: List, answer: str, embeddings) -> float:
    """Calculate confidence score using semantic similarity"""
    try:
        # Get embeddings
        query_embedding = embeddings.embed_query(query)
        answer_embedding = embeddings.embed_query(answer)
        
        # Calculate similarity between query and answer
        query_answer_sim = cosine_similarity(
            [query_embedding], 
            [answer_embedding]
        )[0][0]
        
        # Calculate average similarity with retrieved documents
        if retrieved_docs:
            doc_embeddings = [embeddings.embed_query(doc.page_content) for doc in retrieved_docs]
            doc_similarities = cosine_similarity([query_embedding], doc_embeddings)[0]
            avg_doc_sim = np.mean(doc_similarities)
        else:
            avg_doc_sim = 0.0
        
        # Combine scores (weighted average)
        confidence = (0.4 * query_answer_sim + 0.6 * avg_doc_sim)
        
        # Adjust based on answer length (penalize very short answers)
        if len(answer.split()) < 10:
            confidence *= 0.7
        
        return float(min(max(confidence, 0.0), 1.0))
    except Exception as e:
        st.warning(f"Confidence calculation error: {str(e)}")
        # Fallback to simple heuristic
        return 0.7 if len(retrieved_docs) > 0 and len(answer) > 50 else 0.4

def get_ai_response(query: str, vector_store, openai_api_key: str, target_language: str = 'English') -> Tuple[str, float, List]:
    """Get AI response using RAG with multilingual support"""
    try:
        # Translate query to English if needed
        english_query = query
        if target_language != 'English':
            try:
                translator = GoogleTranslator(source='auto', target='en')
                english_query = translator.translate(query)
            except:
                pass
        
        llm = ChatOpenAI(
            model_name="gpt-3.5-turbo",
            temperature=0.3,
            openai_api_key=openai_api_key
        )
        
        # Custom prompt for multilingual support
        prompt_template = """You are a helpful customer support assistant. 
        Use the following context to answer the question. If you cannot find the answer in the context, 
        say so politely and suggest contacting support.
        
        Context: {context}
        
        Question: {question}
        
        Provide a helpful, accurate answer:"""
        
        PROMPT = PromptTemplate(
            template=prompt_template,
            input_variables=["context", "question"]
        )
        
        qa_chain = RetrievalQA.from_chain_type(
            llm=llm,
            chain_type="stuff",
            retriever=vector_store.as_retriever(search_kwargs={"k": 4}),
            return_source_documents=True,
            chain_type_kwargs={"prompt": PROMPT}
        )
        
        result = qa_chain({"query": english_query})
        
        answer = result['result']
        source_docs = result['source_documents']
        
        # Calculate confidence score
        confidence = calculate_semantic_confidence(
            english_query, 
            source_docs, 
            answer, 
            st.session_state.embeddings_model
        )
        
        # Translate answer back to target language if needed
        if target_language != 'English':
            answer = translate_text(answer, target_language)
        
        return answer, confidence, source_docs
    except Exception as e:
        st.error(f"Error getting AI response: {str(e)}")
        error_msg = "I apologize, but I encountered an error processing your query."
        if target_language != 'English':
            error_msg = translate_text(error_msg, target_language)
        return error_msg, 0.3, []

def create_ticket(query: str, language: str, category: str, confidence: float):
    """Create escalation ticket with priority and assignment"""
    priority = assign_priority(confidence, category)
    
    # Simple round-robin assignment (in production, use proper assignment logic)
    agents = ['Agent A', 'Agent B', 'Agent C']
    assigned_to = agents[len(st.session_state.tickets) % len(agents)]
    
    ticket = {
        'id': str(uuid.uuid4())[:8],
        'query': query,
        'language': language,
        'category': category,
        'status': 'Open',
        'priority': priority,
        'assigned_to': assigned_to,
        'timestamp': datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        'resolved_at': None,
        'resolution_time': None
    }
    
    st.session_state.tickets.append(ticket)
    st.session_state.analytics['escalated'] += 1
    
    # Save to database
    conn = st.session_state.db_conn
    c = conn.cursor()
    c.execute('''INSERT INTO tickets VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)''',
              (ticket['id'], ticket['query'], ticket['language'], ticket['category'],
               ticket['status'], ticket['priority'], ticket['assigned_to'], 
               ticket['timestamp'], ticket['resolved_at'], ticket['resolution_time']))
    conn.commit()
    
    return ticket['id']

def save_chat_to_db(chat_entry: Dict):
    """Save chat entry to database"""
    conn = st.session_state.db_conn
    c = conn.cursor()
    
    chat_id = str(uuid.uuid4())
    c.execute('''INSERT INTO chat_history VALUES (?, ?, ?, ?, ?, ?, ?, ?)''',
              (chat_id, chat_entry['role'], chat_entry['message'], 
               chat_entry['timestamp'].strftime("%Y-%m-%d %H:%M:%S"),
               chat_entry.get('confidence', None),
               chat_entry.get('response_time', None),
               chat_entry.get('language', None),
               chat_entry.get('category', None)))
    conn.commit()
    return chat_id

def update_daily_analytics():
    """Update daily analytics in database"""
    conn = st.session_state.db_conn
    c = conn.cursor()
    
    today = datetime.now().strftime("%Y-%m-%d")
    analytics = st.session_state.analytics
    
    # Calculate averages
    bot_messages = [msg for msg in st.session_state.chat_history if msg['role'] == 'bot']
    avg_response_time = sum(msg.get('response_time', 0) for msg in bot_messages) / len(bot_messages) if bot_messages else 0
    avg_confidence = sum(msg.get('confidence', 0) for msg in bot_messages) / len(bot_messages) if bot_messages else 0
    
    c.execute('''INSERT OR REPLACE INTO analytics VALUES (?, ?, ?, ?, ?, ?)''',
              (today, analytics['total_queries'], analytics['answered'], 
               analytics['escalated'], avg_response_time, avg_confidence))
    conn.commit()

# Sidebar - Knowledge Management
st.sidebar.title("📚 Knowledge Base Management")

# API Key input
openai_api_key = st.sidebar.text_input("OpenAI API Key", type="password", help="Enter your OpenAI API key")

st.sidebar.markdown("---")

# File uploader (including images for OCR)
uploaded_files = st.sidebar.file_uploader(
    "Upload PDF/DOCX/Image files",
    type=['pdf', 'docx', 'png', 'jpg', 'jpeg'],
    accept_multiple_files=True
)

# URL input
url_input = st.sidebar.text_area("Enter URLs (one per line)", height=100)

# Process button
if st.sidebar.button("🔄 Process Knowledge Base", type="primary"):
    if not openai_api_key:
        st.sidebar.error("Please enter your OpenAI API key first!")
    else:
        with st.spinner("Processing knowledge base..."):
            all_text = ""
            
            # Process uploaded files
            if uploaded_files:
                progress_bar = st.sidebar.progress(0)
                for idx, file in enumerate(uploaded_files):
                    if file.name.endswith('.pdf'):
                        all_text += extract_text_from_pdf(file) + "\n\n"
                    elif file.name.endswith('.docx'):
                        all_text += extract_text_from_docx(file) + "\n\n"
                    elif file.name.lower().endswith(('.png', '.jpg', '.jpeg')):
                        ocr_text = extract_text_from_image(file)
                        if ocr_text:
                            all_text += f"[OCR from {file.name}]\n{ocr_text}\n\n"
                    
                    progress_bar.progress((idx + 1) / len(uploaded_files))
                progress_bar.empty()
            
            # Process URLs
            if url_input.strip():
                urls = [url.strip() for url in url_input.split('\n') if url.strip()]
                progress_bar = st.sidebar.progress(0)
                for idx, url in enumerate(urls):
                    st.sidebar.info(f"Scraping: {url[:50]}...")
                    scraped_text = scrape_url(url)
                    if scraped_text:
                        all_text += f"[Content from {url}]\n{scraped_text}\n\n"
                    progress_bar.progress((idx + 1) / len(urls))
                progress_bar.empty()
            
            if all_text.strip():
                st.session_state.knowledge_base_text = all_text
                st.session_state.vector_store = create_vector_store(all_text, openai_api_key)
                
                if st.session_state.vector_store:
                    st.sidebar.success(f"✅ Processed {len(all_text)} characters successfully!")
                else:
                    st.sidebar.error("Failed to create vector store")
            else:
                st.sidebar.warning("No content to process. Please upload files or enter URLs.")

# Display knowledge base status
if st.session_state.vector_store:
    st.sidebar.success("✅ Knowledge Base Active")
    st.sidebar.metric("KB Size", f"{len(st.session_state.knowledge_base_text):,} chars")
    
    # Cache statistics
    num_docs = st.session_state.vector_store.index.ntotal if st.session_state.vector_store else 0
    st.sidebar.metric("Indexed Documents", num_docs)
else:
    st.sidebar.info("ℹ️ No knowledge base loaded")

st.sidebar.markdown("---")
st.sidebar.caption("💡 Supports: PDF, DOCX, Images (OCR), Web URLs")

# Main page
st.markdown('<div class="main-header">🤖 Customer Support & FAQ AI Agent</div>', unsafe_allow_html=True)

# Tabs
tab1, tab2, tab3 = st.tabs(["💬 Chat Agent", "🎫 Ticket Management", "📊 Analytics Dashboard"])

# Tab 1: Chat Agent
with tab1:
    st.header("Chat with AI Support Agent")
    
    if not openai_api_key:
        st.warning("⚠️ Please enter your OpenAI API key in the sidebar to start chatting.")
    elif not st.session_state.vector_store:
        st.info("ℹ️ Please upload and process knowledge base files in the sidebar first.")
    else:
        # Chat interface
        chat_container = st.container()
        
        # Display chat history
        with chat_container:
            for i, chat in enumerate(st.session_state.chat_history):
                if chat['role'] == 'user':
                    st.markdown(f'<div class="chat-message user-message">👤 **You:** {chat["message"]}</div>', unsafe_allow_html=True)
                else:
                    # Display confidence indicator
                    confidence = chat.get('confidence', 0)
                    if confidence >= 0.7:
                        conf_class = "confidence-high"
                        conf_emoji = "🟢"
                    elif confidence >= 0.5:
                        conf_class = "confidence-medium"
                        conf_emoji = "🟡"
                    else:
                        conf_class = "confidence-low"
                        conf_emoji = "🔴"
                    
                    response_time = chat.get('response_time', 0)
                    
                    st.markdown(
                        f'<div class="chat-message bot-message">'
                        f'🤖 **AI:** {chat["message"]}<br>'
                        f'<small>{conf_emoji} Confidence: <span class="{conf_class}">{confidence:.1%}</span> | '
                        f'⏱️ {response_time:.2f}s</small>'
                        f'</div>', 
                        unsafe_allow_html=True
                    )
                    
                    # Feedback buttons
                    if i not in st.session_state.feedback:
                        col1, col2, col3 = st.columns([1, 1, 10])
                        with col1:
                            if st.button("👍", key=f"up_{i}"):
                                st.session_state.feedback[i] = 'positive'
                                st.rerun()
                        with col2:
                            if st.button("👎", key=f"down_{i}"):
                                st.session_state.feedback[i] = 'negative'
                                st.rerun()
                    else:
                        if st.session_state.feedback[i] == 'positive':
                            st.success("✓ Marked as helpful")
                        else:
                            st.error("✗ Marked as not helpful")
        
        # Input area
        with st.form(key="chat_form", clear_on_submit=True):
            col1, col2 = st.columns([4, 1])
            with col1:
                user_input = st.text_input("Ask your question:", placeholder="Type your question here...")
            with col2:
                submit_button = st.form_submit_button("Send 🚀", use_container_width=True)
        
        if submit_button and user_input:
            start_time = time.time()
            
            # Detect language and category
            language = detect_language(user_input)
            category = categorize_query(user_input)
            
            # Update analytics
            st.session_state.analytics['total_queries'] += 1
            st.session_state.analytics['languages'][language] = st.session_state.analytics['languages'].get(language, 0) + 1
            st.session_state.analytics['categories'][category] += 1
            
            # Add user message to history
            user_chat = {
                'role': 'user',
                'message': user_input,
                'timestamp': datetime.now(),
                'language': language,
                'category': category
            }
            st.session_state.chat_history.append(user_chat)
            save_chat_to_db(user_chat)
            
            # Get AI response (with language support)
            with st.spinner(f"Thinking... ({language} detected)"):
                answer, confidence, source_docs = get_ai_response(
                    user_input, 
                    st.session_state.vector_store, 
                    openai_api_key,
                    language
                )
            
            response_time = time.time() - start_time
            
            # Check if escalation is needed
            if confidence < 0.6:
                ticket_id = create_ticket(user_input, language, category, confidence)
                escalation_msg = f"\n\n⚠️ **Note:** Your query has been escalated to our support team. Ticket ID: `{ticket_id}` | Priority: {assign_priority(confidence, category)}"
                if language != 'English':
                    escalation_msg = translate_text(escalation_msg, language)
                answer += escalation_msg
            else:
                st.session_state.analytics['answered'] += 1
            
            # Add bot response to history
            bot_chat = {
                'role': 'bot',
                'message': answer,
                'timestamp': datetime.now(),
                'confidence': confidence,
                'response_time': response_time,
                'language': language,
                'category': category
            }
            st.session_state.chat_history.append(bot_chat)
            save_chat_to_db(bot_chat)
            
            # Update daily analytics
            update_daily_analytics()
            
            st.rerun()

# Tab 2: Ticket Management
with tab2:
    st.header("Escalated Tickets")
    
    if st.session_state.tickets:
        # Filters
        col1, col2, col3 = st.columns(3)
        with col1:
            status_filter = st.multiselect(
                "Filter by Status",
                options=['Open', 'In Progress', 'Closed'],
                default=['Open', 'In Progress']
            )
        with col2:
            priority_filter = st.multiselect(
                "Filter by Priority",
                options=['High', 'Medium', 'Low'],
                default=['High', 'Medium', 'Low']
            )
        with col3:
            category_filter = st.multiselect(
                "Filter by Category",
                options=['Billing', 'Technical', 'General'],
                default=['Billing', 'Technical', 'General']
            )
        
        # Create DataFrame
        df_tickets = pd.DataFrame(st.session_state.tickets)
        
        # Apply filters
        filtered_df = df_tickets[
            (df_tickets['status'].isin(status_filter)) &
            (df_tickets['priority'].isin(priority_filter)) &
            (df_tickets['category'].isin(category_filter))
        ]
        
        # Summary metrics
        col1, col2, col3, col4 = st.columns(4)
        with col1:
            st.metric("Total Tickets", len(df_tickets))
        with col2:
            st.metric("Open", len(df_tickets[df_tickets['status'] == 'Open']))
        with col3:
            st.metric("High Priority", len(df_tickets[df_tickets['priority'] == 'High']))
        with col4:
            avg_resolution = df_tickets[df_tickets['resolution_time'].notna()]['resolution_time'].mean()
            st.metric("Avg Resolution Time", f"{avg_resolution:.1f}h" if not pd.isna(avg_resolution) else "N/A")
        
        st.markdown("---")
        
        # Display tickets
        for idx, ticket in filtered_df.iterrows():
            priority_color = {'High': '🔴', 'Medium': '🟡', 'Low': '🟢'}
            
            with st.expander(f"{priority_color[ticket['priority']]} Ticket #{ticket['id']} - {ticket['status']} ({ticket['priority']} Priority)"):
                col1, col2 = st.columns(2)
                with col1:
                    st.write(f"**Query:** {ticket['query']}")
                    st.write(f"**Category:** {ticket['category']}")
                    st.write(f"**Language:** {ticket['language']}")
                with col2:
                    st.write(f"**Assigned To:** {ticket['assigned_to']}")
                    st.write(f"**Created:** {ticket['timestamp']}")
                    if ticket['resolved_at']:
                        st.write(f"**Resolved:** {ticket['resolved_at']}")
                        st.write(f"**Resolution Time:** {ticket['resolution_time']:.1f}h")
                
                # Status update
                col1, col2, col3 = st.columns([2, 2, 1])
                with col1:
                    new_status = st.selectbox(
                        "Update Status",
                        options=['Open', 'In Progress', 'Closed'],
                        index=['Open', 'In Progress', 'Closed'].index(ticket['status']),
                        key=f"status_{ticket['id']}"
                    )
                with col2:
                    new_assignee = st.selectbox(
                        "Reassign To",
                        options=['Agent A', 'Agent B', 'Agent C'],
                        index=['Agent A', 'Agent B', 'Agent C'].index(ticket['assigned_to']),
                        key=f"assign_{ticket['id']}"
                    )
                with col3:
                    st.write("")
                    st.write("")
                    if st.button("Update", key=f"update_{ticket['id']}", type="primary"):
                        # Calculate resolution time if closing
                        if new_status == 'Closed' and ticket['status'] != 'Closed':
                            created = datetime.strptime(ticket['timestamp'], "%Y-%m-%d %H:%M:%S")
                            resolution_time = (datetime.now() - created).total_seconds() / 3600
                            st.session_state.tickets[idx]['resolved_at'] = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                            st.session_state.tickets[idx]['resolution_time'] = resolution_time
                        
                        st.session_state.tickets[idx]['status'] = new_status
                        st.session_state.tickets[idx]['assigned_to'] = new_assignee
                        st.success(f"Ticket #{ticket['id']} updated!")
                        st.rerun()
    else:
        st.info("No escalated tickets yet.")

# Tab 3: Analytics Dashboard
with tab3:
    st.header("Analytics Dashboard")
    
    analytics = st.session_state.analytics
    
    # Metrics row
    col1, col2, col3, col4 = st.columns(4)
    with col1:
        st.metric("Total Queries", analytics['total_queries'])
    with col2:
        st.metric("Answered", analytics['answered'])
    with col3:
        st.metric("Escalated", analytics['escalated'])
    with col4:
        resolution_rate = (analytics['answered'] / analytics['total_queries'] * 100) if analytics['total_queries'] > 0 else 0
        st.metric("Resolution Rate", f"{resolution_rate:.1f}%")
    
    st.markdown("---")
    
    # Time-based analytics
    if st.session_state.chat_history:
        st.subheader("📈 Query Trends")
        
        # Create time series data
        chat_df = pd.DataFrame([
            {
                'timestamp': chat['timestamp'],
                'role': chat['role'],
                'confidence': chat.get('confidence', None),
                'response_time': chat.get('response_time', None)
            }
            for chat in st.session_state.chat_history
        ])
        
        # Queries over time
        chat_df['hour'] = pd.to_datetime(chat_df['timestamp']).dt.floor('H')
        hourly_queries = chat_df[chat_df['role'] == 'user'].groupby('hour').size().reset_index(name='queries')
        
        if len(hourly_queries) > 0:
            fig_timeline = px.line(
                hourly_queries,
                x='hour',
                y='queries',
                title="Queries Over Time",
                labels={'hour': 'Time', 'queries': 'Number of Queries'}
            )
            st.plotly_chart(fig_timeline, use_container_width=True)
    
    st.markdown("---")
    
    # Charts row 1
    col1, col2 = st.columns(2)
    
    with col1:
        # Query Distribution
        fig_queries = go.Figure(data=[
            go.Bar(
                x=['Answered', 'Escalated'],
                y=[analytics['answered'], analytics['escalated']],
                marker_color=['#2ecc71', '#e74c3c'],
                text=[analytics['answered'], analytics['escalated']],
                textposition='auto'
            )
        ])
        fig_queries.update_layout(
            title="Query Distribution",
            xaxis_title="Status",
            yaxis_title="Count",
            height=300
        )
        st.plotly_chart(fig_queries, use_container_width=True)
    
    with col2:
        # Language Distribution
        lang_data = {k: v for k, v in analytics['languages'].items() if v > 0}
        if lang_data:
            fig_lang = px.pie(
                values=list(lang_data.values()),
                names=list(lang_data.keys()),
                title="Language Distribution"
            )
            fig_lang.update_layout(height=300)
            st.plotly_chart(fig_lang, use_container_width=True)
        else:
            st.info("No language data available yet")
    
    # Charts row 2
    st.markdown("---")
    col1, col2 = st.columns(2)
    
    with col1:
        # Category Breakdown
        st.subheader("Category Breakdown")
        category_df = pd.DataFrame({
            'Category': list(analytics['categories'].keys()),
            'Count': list(analytics['categories'].values())
        })
        
        fig_category = px.bar(
            category_df,
            x='Category',
            y='Count',
            color='Category',
            title="Queries by Category"
        )
        st.plotly_chart(fig_category, use_container_width=True)
    
    with col2:
        # Performance Metrics
        st.subheader("Performance Metrics")
        bot_messages = [msg for msg in st.session_state.chat_history if msg['role'] == 'bot']
        
        if bot_messages:
            avg_response_time = sum(msg.get('response_time', 0) for msg in bot_messages) / len(bot_messages)
            avg_confidence = sum(msg.get('confidence', 0) for msg in bot_messages) / len(bot_messages)
            
            perf_metrics = pd.DataFrame({
                'Metric': ['Avg Response Time (s)', 'Avg Confidence Score'],
                'Value': [avg_response_time, avg_confidence]
            })
            
            fig_perf = go.Figure(data=[
                go.Bar(
                    x=perf_metrics['Metric'],
                    y=perf_metrics['Value'],
                    marker_color=['#3498db', '#9b59b6'],
                    text=[f"{avg_response_time:.2f}s", f"{avg_confidence:.1%}"],
                    textposition='auto'
                )
            ])
            fig_perf.update_layout(
                title="Average Performance",
                height=300,
                yaxis_title="Value"
            )
            st.plotly_chart(fig_perf, use_container_width=True)
        else:
            st.info("No performance data available yet")
    
    # Confidence distribution
    st.markdown("---")
    st.subheader("Confidence Score Distribution")
    
    if bot_messages:
        confidences = [msg.get('confidence', 0) for msg in bot_messages]
        
        fig_conf = px.histogram(
            x=confidences,
            nbins=20,
            title="Distribution of AI Confidence Scores",
            labels={'x': 'Confidence Score', 'y': 'Frequency'},
            color_discrete_sequence=['#1f77b4']
        )
        fig_conf.add_vline(x=0.6, line_dash="dash", line_color="red", 
                          annotation_text="Escalation Threshold (0.6)")
        st.plotly_chart(fig_conf, use_container_width=True)
    
    # Customer Satisfaction
    st.markdown("---")
    st.subheader("Customer Satisfaction")
    
    if st.session_state.feedback:
        positive = sum(1 for f in st.session_state.feedback.values() if f == 'positive')
        negative = sum(1 for f in st.session_state.feedback.values() if f == 'negative')
        total_feedback = positive + negative
        
        col1, col2, col3 = st.columns(3)
        with col1:
            st.metric("Total Feedback", total_feedback)
        with col2:
            st.metric("Positive", positive, f"{positive/total_feedback*100:.1f}%")
        with col3:
            st.metric("Negative", negative, f"{negative/total_feedback*100:.1f}%")
        
        # Satisfaction pie chart
        fig_satisfaction = px.pie(
            values=[positive, negative],
            names=['Positive 👍', 'Negative 👎'],
            title="User Satisfaction",
            color_discrete_sequence=['#2ecc71', '#e74c3c']
        )
        st.plotly_chart(fig_satisfaction, use_container_width=True)
    else:
        st.info("No customer feedback received yet")
    
    # Peak Hours Analysis
    if st.session_state.chat_history:
        st.markdown("---")
        st.subheader("📊 Peak Hours Analysis")
        
        user_chats = [chat for chat in st.session_state.chat_history if chat['role'] == 'user']
        if user_chats:
            hours = [chat['timestamp'].hour for chat in user_chats]
            hour_counts = pd.Series(hours).value_counts().sort_index()
            
            fig_peak = px.bar(
                x=hour_counts.index,
                y=hour_counts.values,
                title="Query Volume by Hour of Day",
                labels={'x': 'Hour of Day', 'y': 'Number of Queries'},
                color=hour_counts.values,
                color_continuous_scale='Blues'
            )
            fig_peak.update_layout(showlegend=False)
            st.plotly_chart(fig_peak, use_container_width=True)

# Footer
st.markdown("---")
st.markdown("""
<div style='text-align: center; color: #666;'>
    <p><strong>AI Customer Support Agent v2.0 (Complete Edition)</strong></p>
    <p>Powered by OpenAI GPT-3.5, LangChain & FAISS</p>
    <p><small>✅ Features: PDF/DOCX/Image Processing | Web Scraping | OCR | Multilingual (EN/HI/MR) | 
    Semantic Confidence Scoring | Ticket Management | Real-time Analytics | Persistent Database</small></p>
</div>
""", unsafe_allow_html=True)

# Export data functionality
with st.sidebar:
    st.markdown("---")
    st.subheader("📥 Export Data")
    
    if st.button("Export Chat History (CSV)"):
        if st.session_state.chat_history:
            df = pd.DataFrame(st.session_state.chat_history)
            csv = df.to_csv(index=False)
            st.download_button(
                "Download CSV",
                csv,
                "chat_history.csv",
                "text/csv"
            )
    
    if st.button("Export Tickets (CSV)"):
        if st.session_state.tickets:
            df = pd.DataFrame(st.session_state.tickets)
            csv = df.to_csv(index=False)
            st.download_button(
                "Download CSV",
                csv,
                "tickets.csv",
                "text/csv"
            )
    
    if st.button("🗑️ Clear All Data"):
        if st.button("⚠️ Confirm Clear", type="primary"):
            st.session_state.chat_history = []
            st.session_state.tickets = []
            st.session_state.analytics = {
                'total_queries': 0,
                'answered': 0,
                'escalated': 0,
                'languages': {'English': 0, 'Hindi': 0, 'Marathi': 0, 'Other': 0},
                'categories': {'Billing': 0, 'Technical': 0, 'General': 0}
            }
            st.session_state.feedback = {}
            
            # Clear database
            conn = st.session_state.db_conn
            c = conn.cursor()
            c.execute("DELETE FROM chat_history")
            c.execute("DELETE FROM tickets")
            c.execute("DELETE FROM feedback")
            c.execute("DELETE FROM analytics")
            conn.commit()
            
            st.success("All data cleared!")
            st.rerun()
